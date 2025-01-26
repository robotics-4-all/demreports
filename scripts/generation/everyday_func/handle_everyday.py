"""
Function that handles the everyday functionality section of the report
"""
# pylint: disable=C0301

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from generation.utilities.create_literal_list import create_literal_list
from generation.tests_literals.fucas_literals import fucas_literals
from generation.tests_literals.frssd_literals import frssd_literals

def handle_everyday_functionality(parsed, document, literals, print_output = False):
    """
    Handles the functionality for generating a report based on everyday activities.

    Args:
        parsed (dict): A dictionary containing parsed data from the input.
        document (Document): A Document object where the report will be added.
        literals (dict): A dictionary containing literal strings used in the report.
        print_output (bool, optional): If True, prints the generated report text. Defaults to False.

    Returns:
        None
    """
    p1 = document.add_paragraph()

    # obj: 7: no, >=9 [ήπια, σοβαρά] ελλείματα
    # obj_medication: λήψη της φαρμακευτικής αγωγής
    # obj_telephone: επικοινωνίας με την χρήση τηλεφώνου
    # obj_financial: ικανότητα οικονομικών συναλλαγών
    # obj_hygene: διατήρησης της προσωπικής υγιεινής
    # obj_orientation: προσανατολισμού σε χώρο
    # obj_dressing: ένδυσης

    fucas_lits = fucas_literals(parsed['fucas'])
    frssd_lits = frssd_literals(parsed['frssd'])

    fucas_admin = parsed['fucas'].administered
    frssd_admin = parsed['frssd'].administered

    printable_1 = "Από τα αποτελέσματα της αντικειμενικής εκτίμησης μέσω της δοκιμασίας καθημερινής λειτουργικότητας (FUCAS)"
    if len(fucas_lits['objective']['no']) > 0:
        printable_1 += f" δεν διαπιστώθηκαν ελλείματα {'στην ικανότητα' if len(fucas_lits['objective']['no']) == 1 else 'στις ικανότητες'} {create_literal_list(fucas_lits['objective']['no'])}"
    if len(fucas_lits['objective']['yes']) > 0:
        intermed = ", ενώ" if len(fucas_lits['objective']['no']) > 0 else ""
        printable_1 += f"{intermed} διαπιστώθηκαν ελλείματα {'στην ικανότητα' if len(fucas_lits['objective']['yes']) == 1 else 'στις ικανότητες'} {create_literal_list(fucas_lits['objective']['yes'])}"
    printable_1 += ". "

    if fucas_admin:
        p1.add_run(printable_1)
        if print_output:
            print(printable_1)

    # 0: no, 1: mild, 2: μέτρια, 3: σοβαρά
    # nutrition: διατροφής
    # dressing: ένδυσης
    # akrateia: αναφέρθηκε εγκράτεια ούρων (3+κοπράνων)
    #== check presentation

    atts_exist = parsed["patient"].att_1_existed or parsed["patient"].att_2_existed
    atts_str = f"με {literals['att_literal_1']} "

    printable_2 = f"Σύμφωνα με την συνέντευξη που πραγματοποιήθηκε {atts_str} και μετά από την χορήγηση ημι-δομημένου ερωτηματολογίου (FRSSD){frssd_lits['finals']['mild']}{frssd_lits['finals']['moderate']}{frssd_lits['finals']['severe']}."

    if frssd_admin:
        p1.add_run(printable_2)

    printable_3 = f" Τα παραπάνω ευρήματα συνηγορούν στο ότι για το χρονικό διάστημα στο οποίο αναφέρεται η νευροψυχολογική εκτίμηση, {literals['article_lastname_mr_mrs']} χρειαζόταν υπενθύμιση και βοήθεια μέσω τρίτων προσώπων προκειμένου να μπορεί να ανταπεξέλθει στις σύνθετες δραστηριότητες της καθημερινής ζωής."
    if frssd_admin and (frssd_lits['finals']['moderate'] != "" or frssd_lits['finals']['severe'] != ""):
        p1.add_run(printable_3)

    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
