# pylint: disable=C0301
"""
File that contains the function to handle the generation of the executive functions section in a document.
"""

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from generation.tests_literals.fucas_literals import fucas_literals
from generation.utilities.create_literal_list import create_literal_list

def handle_executive_functions(parsed, document, literals):
    """
    Handles the generation of the executive functions section in a document.

    Args:
        parsed (dict): A dictionary containing parsed data, including 'fucas' key.
        document (Document): A python-docx Document object where the content will be added.
        literals (dict): A dictionary containing literal strings used in the document.
        print_output (bool, optional): If True, prints the output. Defaults to False.

    Returns:
        None
    """
    p1 = document.add_paragraph()

    fl = fucas_literals(parsed['fucas'])

    # input from Elena
    porisma = ""
    negs = ""
    post = ""
    if len(fl['individual_affected']) > 0:
        negs = f"διαπιστώθηκαν ελλείψεις στις ικανότητες {fl['affected']}"
        porisma += negs
    if len(fl['individual_unaffected']) > 0:
        post = f"δεν διαπιστώθηκαν ελλείψεις στις ικανότητες {fl['unaffected']}"
        if len(fl['individual_affected']) > 0:
            post = f". Σε αντίθεση, {post}"
        porisma += post

    printable_1 = f"Όσον αφορά την αξιολόγηση των σύνθετων και των απλών καθημερινών δραστηριοτήτων, η οποία έγινε μέσω της δοκιμασίας FUCAS, {porisma}. "
    p1.add_run(printable_1)

    p1.add_run(f"Τα παραπάνω ευρήματα καταδεικνύουν πως {literals['article_lastname_mr_mrs']}")
    p1.add_run(f"{'δεν' if len(fl['individual_affected']) == 0 else ''} παρουσίαζε δυσκολίες σε ικανότητες που είναι απαραίτητες προκειμένου να ολοκληρωθεί σωστά η εκτέλεση σύνθετων νοητικών έργων/δραστηριοτήτων.")

    if len(fl['individual_affected']) > 0:
        p1.add_run(f" Πιθανά παραδείγματα που να σχετίζονται με την καθημερινή ζωή αφορούν τη δυσκολία {literals['examinee_gender']} ")
        lst = create_literal_list([f"{fl['examples'][item]}" for item in fl['individual_affected']], add_identifiers=True)
        p1.add_run(lst)
    p1.add_run(".")

    if "ακρίβειας βημάτων για την διεκπεραίωση έργων" in fl['individual_affected'] and "προοπτικής μνήμης" in fl['individual_affected']:
        p1.add_run(" Διαπιστώθηκε επίσης ότι όταν το σύνθετο έργο που πρέπει να ολοκληρωθεί έχει απαιτήσεις σε ικανότητες μακρόχρονης μνήμης, εμφανίζονται δυσκολίες σε επίπεδο ακρίβειας βημάτων οι οποίες πρέπει να ακολουθηθούν από τον εξεταζόμενο προκειμένου να διεκπεραιωθεί το έργο.")

    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
