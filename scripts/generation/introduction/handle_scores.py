# pylint: disable=C0301
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt

from generation.tests_literals.npi_literals import npi_literals
from generation.tests_literals.ravlt_literals import ravlt_literals
from generation.tests_literals.rocft_literals import rocft_literals

def handle_scores(parsed, document):
    """
    Processes and formats various cognitive and psychological test scores from the parsed data 
    and adds them as a paragraph to the provided document.

    Args:
        parsed (dict): A dictionary containing the test scores. The keys are the test names 
                       (e.g., 'mmse', 'ravlt', 'rocft', etc.) and the values are objects with 
                       a 'score' attribute or other relevant attributes.
        document (Document): A document object where the formatted scores will be added as a paragraph.

    Returns:
        None
    """
    npi_lits = npi_literals(parsed['npi'])
    ravlt_lits = ravlt_literals(parsed['ravlt'], parsed['patient'].age)
    rocft_lits = rocft_literals(parsed['rocft'], parsed['patient'].age, parsed['patient'].education)

    p1 = document.add_paragraph()

    if 'mmse' in parsed and parsed['mmse'].administered:
        mmse_lit = 'MMSE' if parsed['patient'].education > 4 else 'HMSE'
        p1.add_run(f"{mmse_lit}: ")
        if parsed['mmse'].score < 24:
            p1.add_run(str(parsed['mmse'].score)).bold = True
        else:
            p1.add_run(f"{str(parsed['mmse'].score)}")
        p1.add_run("/30\n")

    if 'ravlt' in parsed and parsed['ravlt'].administered:
        p1.add_run("RAVLT Ικανότητα μάθησης / Ικανότητα συγκράτησης: ")
        if ravlt_lits['learning_problem']:
            p1.add_run(str(parsed['ravlt'].l_score)).bold = True
        else:
            p1.add_run(str(parsed['ravlt'].l_score))

        p1.add_run("/")

        if ravlt_lits['maintain_problem']:
            p1.add_run(str(parsed['ravlt'].m_score)).bold = True
        else:
            p1.add_run(str(parsed['ravlt'].m_score))
        p1.add_run("/15\n")

    if 'rocft' in parsed and parsed['rocft'].administered:
        p1.add_run("ROCFT: Αντιγ/Καθυστ.Ανάκλ.: ")
        if rocft_lits['copy_flag']:
            p1.add_run(str(parsed['rocft'].c_score)).bold = True
        else:
            p1.add_run(str(parsed['rocft'].c_score))

        p1.add_run("/")

        if rocft_lits['recall_problem']:
            p1.add_run(str(parsed['rocft'].d_score)).bold = True
        else:
            p1.add_run(str(parsed['rocft'].d_score))
        p1.add_run("/36\n")

    if 'fucas' in parsed and parsed['fucas'].administered:
        p1.add_run("FUCAS: ")
        if parsed['fucas'].score > 42:
            p1.add_run(f"{str(parsed['fucas'].score)}").bold = True
        else:
            p1.add_run(f"{str(parsed['fucas'].score)}")
        p1.add_run("/126\n")

    if 'frssd' in parsed and parsed['frssd'].administered:
        p1.add_run("FRSSD: ")
        if parsed['frssd'].score > 5:
            p1.add_run(f"{str(parsed['frssd'].score)}").bold = True
        else:
            p1.add_run(f"{str(parsed['frssd'].score)}")
        p1.add_run("/42\n")

    if 'npi' in parsed and parsed['npi'].administered:
        p1.add_run("NPI: ")
        p1.add_run(f"{str(parsed['npi'].score)}")
        p1.add_run("/120\n")

    if 'gds' in parsed and parsed['gds'].administered:
        p1.add_run("GDS: ")
        if parsed['gds'].score > 5:
            p1.add_run(f"{str(parsed['gds'].score)}").bold = True
        else:
            p1.add_run(f"{str(parsed['gds'].score)}")
        p1.add_run("/15\n")

    if 'sast' in parsed and parsed['sast'].administered:
        p1.add_run("SAST Άγχος: ")
        if parsed['sast'].score > 24:
            p1.add_run(f"{str(parsed['sast'].score)}").bold = True
        else:
            p1.add_run(f"{str(parsed['sast'].score)}")
        p1.add_run("/40\n")

    if 'bdi' in parsed and parsed['bdi'].administered:
        p1.add_run("BDI: ")
        if parsed['bdi'].score > 8:
            p1.add_run(f"{str(parsed['bdi'].score)}").bold = True
        else:
            p1.add_run(f"{str(parsed['bdi'].score)}")
        p1.add_run("/69\n")

    if 'bai' in parsed and parsed['bai'].administered:
        p1.add_run("BAI: ")
        if parsed['bai'].score > 7:
            p1.add_run(f"{str(parsed['bai'].score)}").bold = True
        else:
            p1.add_run(f"{str(parsed['bai'].score)}")
        p1.add_run("/63\n")

    p2 = document.add_paragraph("*Οι τιμές που αναγράφονται με έντονη επισήμανση αφορούν τις ελλειμματικές επιδόσεις σύμφωνα με τα όρια κατωφλίου για τον ελληνικό πληθυσμό σύμφωνα με την βιβλιογραφία που παρατίθεται στο τέλος της παρούσας αναφοράς.")
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p2.runs[0].font.size = Pt(8)
