"""
File that handles the generation of the verbal memory paragraph
"""
# pylint: disable=C0301

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from generation.tests_literals.ravlt_literals import ravlt_literals

def handle_verbal_memory(parsed, document, literals, print_output = False):
    """
    Handles the generation of a paragraph related to verbal memory (RAVLT) in a document.

    Args:
        parsed (dict): A dictionary containing parsed data, including 'ravlt' and 'patient' information.
        document (Document): A python-docx Document object where the paragraph will be added.
        literals (dict): A dictionary containing literal strings used in the paragraph.
        print_output (bool, optional): If True, the generated paragraph text will be printed to the console. Defaults to False.

    Returns:
        None
    """
    # check we have RAVLT data
    if 'ravlt' not in parsed or parsed['ravlt'].administered == False:
        return

    ravlt_lits = ravlt_literals(parsed['ravlt'], parsed['patient'].age)
    p1 = document.add_paragraph()

    p1_r1 = p1.add_run("\nΛεκτική μνήμη επεισοδίων: ")
    p1_r1.italic = True

    d = parsed['patient'].date_npse
    if parsed['patient'].revisit:
        d = parsed['patient'].revisit_date

    printable = f"Κατά τη νευροψυχολογική εκτίμηση στις {d} και συγκεκριμένα μέσω της χορήγησης της δοκιμασίας RAVLT, {literals['article_lastname_mr_mrs']} {ravlt_lits['learning']} στην ικανότητα μάθησης καταλόγου λέξεων μετά από επανάληψη. Το παραπάνω γεγονός καταδείκνυε ότι για το διάστημα στο οποίο έγινε η εκτίμηση, {ravlt_lits['learning_explanation']}. Η ικανότητα ανάσυρσης της πληροφορίας από την μακρόχρονη μνήμη, όπως διαπιστώθηκε από την ίδια δοκιμασία, {ravlt_lits['maintain']}, καθώς {literals['article_lastname_mr_mrs']} {ravlt_lits['maintain_explanation']}{ravlt_lits['apospasmatika']}. "

    p1.add_run(printable)

    if print_output:
        print(printable)

    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
