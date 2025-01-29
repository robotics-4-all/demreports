"""
File that handles the generation of the visual memory paragraph in the report
"""
# pylint: disable=C0301

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from generation.tests_literals.rocft_literals import rocft_literals

def handle_visual_memory(parsed, document, literals, print_output = False):
    """
    Processes and adds a paragraph to the document based on ROCFT data from the parsed input.

    Args:
        parsed (dict): A dictionary containing parsed data, including 'rocft' 
        and 'patient' information.
        document (Document): A python-docx Document object where the paragraph will be added.
        literals (dict): A dictionary containing literal strings used in the paragraph.
        print_output (bool, optional): If True, the generated paragraph text will be
        printed to the console. Defaults to False.

    Returns:
        None
    """
    # check we have ROCFT data
    if 'rocft' not in parsed or parsed['rocft'].administered == False:
        return

    rocft_lits = rocft_literals(parsed['rocft'], parsed['patient'].age, parsed['patient'].education)
    p1 = document.add_paragraph()

    p1_r1 = p1.add_run("\nΟπτική μνήμη επεισοδίων: ")
    p1_r1.italic = True

    printable = f"Η οπτική μνήμη επεισοδίων {rocft_lits['recall']}, έτσι όπως διαπιστώθηκε από την ανάκληση της σύνθετης φιγούρας ROCFT, αναλόγως της ηλικίας και της εκπαίδευσης {literals['examinee_gender']}. {literals['article_lastname_mr_mrs_capital']} {rocft_lits['recall_explanation']} της φιγούρας, την οποία είχε προηγουμένως κληθεί να αντιγράψει{rocft_lits['apospasmatika']}."

    p1.add_run(printable)

    if print_output:
        print(printable)

    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
