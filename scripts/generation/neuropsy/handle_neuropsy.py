"""
File handling the generation of the neuropsychological symptoms report paragraph
"""
# pylint: disable=C0301
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from generation.tests_literals.gds_literals import gds_literals
from generation.tests_literals.npi_literals import npi_literals
from generation.utilities.create_literal_list import create_literal_list

def handle_neuropsychological_symptoms(parsed, document, literals):
    """
    Handles the generation of a neuropsychological symptoms report paragraph in a document.
    Args:
        parsed (dict): A dictionary containing parsed data, including 'gds' key for GDS literals.
        document (Document): A python-docx Document object where the paragraph will be added.
        literals (dict): A dictionary containing various literals used in the report, such as 'examinee_gender_v2', 'att_literal_1', and 'att_names_with_relations'.
        print_output (bool, optional): If True, prints the output. Defaults to False.
    Returns:
        None
    """
    p1 = document.add_paragraph()

    anxiety = False
    anxiety_status = ""
    questionnaires = []
    if parsed['sast'].administered:
        questionnaires.append("SAST")
    if parsed['sast'].administered and parsed['sast'].score >= 24:
        anxiety = True
        anxiety_status = "διαπιστώθηκε αγχώδης διαταραχή"

    if parsed['bai'].administered:
        questionnaires.append("BAI")
    if parsed['bai'].administered and parsed['bai'].score >= 8:
        anxiety = True
        anxiety_severity = "ήπια"
        if parsed['bai'].score >= 16:
            anxiety_severity = "μέτρια"
        if parsed['bai'].score >= 26:
            anxiety_severity = "σοβαρή"
        anxiety_status = f"διαπιστώθηκε {anxiety_severity} αγχώδης διαταραχή"

    depression = False
    depression_status = ""
    if parsed['bdi'].administered:
        questionnaires.append("BDI")
    if parsed['bdi'].administered and parsed['bdi'].score >= 9:
        depression = True
        depression_severity = "ήπια"
        if parsed['bdi'].score >= 17:
            depression_severity = "μέτρια"
        if parsed['bdi'].score >= 27:
            depression_severity = "σοβαρή"
        depression_status = f"διαπιστώθηκε {depression_severity} κατάθλιψη"

    if parsed['gds'].administered:
        questionnaires.append("GDS")
    if parsed['gds'].administered and parsed['gds'].score >= 6:
        depression = True
        depression_status = gds_literals(parsed['gds'])

    total_anxiety = "δεν διαπιστώθηκε αγχώδης διαταραχή" if not anxiety else anxiety_status
    total_depression = "δεν διαπιστώθηκε κατάθλιψη" if not depression else depression_status

    quests = ", ".join(questionnaires)

    p1.add_run(f"Σύμφωνα με τα ερωτηματολόγια αυτοαναφοράς ({quests}) που χορηγήθηκαν {literals['examinee_gender_v2']}, για την περίοδο που έγινε η εκτίμηση, {total_anxiety} και {total_depression}.")

    npi_lits = npi_literals(parsed['npi'])

    atts_related = f"{literals['att_literal_1']} ({literals['att_names_with_relations']})" if literals['att_names_with_relations'] != "" else literals['the_same_v3']

    p1.add_run(f" Σύμφωνα με {atts_related}")
    if len(npi_lits[3]) > 0:
        p1.add_run(" αναφέρθηκαν συμπεριφορές")
        p1.add_run(" μεγάλης σοβαρότητας").bold = True
        p1.add_run(" όπως ")
        p1.add_run(create_literal_list(npi_lits[3]))
        p1.add_run(". ")
    if len(npi_lits[2]) > 0:
        is_first = len(npi_lits[3]) == 0
        is_last = len(npi_lits[1]) == 0
        ll = ""
        if is_last and not is_first:
            ll = 'Τέλος'
        if not is_last and not is_first:
            ll = 'Επίσης'

        p1.add_run(f"{ll} αναφέρθηκαν συμπεριφορές")
        p1.add_run(" μέτριας σοβαρότητας").bold = True
        p1.add_run(" όπως ")
        p1.add_run(create_literal_list(npi_lits[2]))
        p1.add_run(". ")
    if len(npi_lits[1]) > 0:
        is_only = len(npi_lits[3]) == 0 or len(npi_lits[2]) == 0
        ll = ""
        if not is_only:
            ll = 'Τέλος'
        p1.add_run(f"{ll} αναφέρθηκαν συμπεριφορές")
        p1.add_run(" ήπιας σοβαρότητας").bold = True
        p1.add_run(" όπως ")
        p1.add_run(create_literal_list(npi_lits[1]))
        p1.add_run(". ")

    if len(npi_lits[1]) == 0 and len(npi_lits[2]) == 0 and len(npi_lits[3]) == 0:
        p1.add_run(" δεν αναφέρθηκαν ").bold = True
        p1.add_run(f"νευροψυχιατρικά συμπτώματα σχετικά με {npi_lits['all']}. ")

    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
