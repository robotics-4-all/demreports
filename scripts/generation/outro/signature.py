from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
# pylint: disable=C0301

def handle_signature(document):
    """
    Adds a signature section to the given document.

    This function creates a new paragraph in the document and adds multiple
    lines of text to it, including a signature line and various titles and
    qualifications. The paragraph is then center-aligned.

    Args:
        document (Document): The document object to which the signature
                             section will be added.
    """
    p1 = document.add_paragraph()

    # input from Elena
    p1.add_run("\n\nΥπογραφή\n\n\n\n\n")
    p1.add_run("Πόπτση Ελένη\n")
    p1.add_run("PhD Ψυχολογίας ΑΠΘ\n")
    p1.add_run("Ψυχολόγος ΑΠΘ\n")
    p1.add_run("MSc Κοινωνικής Ψυχιατρικής ΔΠΘ\n")

    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    p2 = document.add_paragraph()
    p2.add_run("Βιβλιογραφικές παραπομπές που αφορούν τα όρια κατωφλίου για τον ελληνικό πληθυσμό:\n\n")
    p2.add_run("1. Fountoulakis, K. N., Tsolaki, M., Chantzi, H., & Kazis, A. (2000). Mini mental state examination (MMSE): a validation study in Greece. American Journal of Alzheimer's Disease & Other Dementias®, 15(6), 342-345.\n")
    p2.add_run("2. Messinis, L., Tsakona, I., Malefaki, S., & Papathanasopoulos, P. (2007). Normative data and discriminant validity of Rey's Verbal Learning Test for the Greek adult population. Archives of Clinical Neuropsychology, 22(6), 739-752.\n")
    p2.add_run("3. Tsatali, M., Emmanouel, A., Gialaouzidis, M., Avdikou, K., Stefanatos, C., Diamantidou, A., ... & Tsolaki, M. (2022). Rey complex figure test (RCFT): Norms for the Greek older adult population. Applied Neuropsychology: Adult, 29(5), 958-966.\n")
    p2.add_run("4. Kounti, F., Tsolaki, M., & Kiosseoglou, G. (2006). Functional cognitive assessment scale (FUCAS): A new scale to assess executive cognitive function in daily life activities in patients with dementia and mild cognitive impairment. Human Psychopharmacology: Clinical and Experimental, 21(5), 305-311.\n")
    p2.add_run("5. Fountoulakis, K. N., Tsolaki, M., Iacovides, A., Yesavage, J., O’Hara, R., Kazis, A., & Ierodiakonou, C. (1999). The validation of the short form of the Geriatric Depression Scale (GDS) in Greece. Aging Clinical and Experimental Research, 11, 367-372. Grammatikopoulos, I. A., Sinoff, G., Alegakis, A., Kounalakis, D., Antonopoulou, M., & Lionis, C. (2010). The short anxiety screening test in Greek: Translation and validation. Annals of General Psychiatry, 9, 1-8.\n")
    p2.add_run("6. Grammatikopoulos, I. A., Sinoff, G., Alegakis, A., Kounalakis, D., Antonopoulou, M., & Lionis, C. (2010). The short anxiety screening test in Greek: Translation and validation. Annals of General Psychiatry, 9, 1-8.\n")
    for run in p2.runs:
        run.font.size = Pt(8)

