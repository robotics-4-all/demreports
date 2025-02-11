"""
File that contains the function to handle the first visit section of the document.
"""

# pylint: disable=C0301
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from generation.utilities.compare_dates import compare_dates

def handle_s1_first_visit(parsed, document, literals):
    """
    Handles the generation of the first visit section for a document.

    This function creates a paragraph in the provided document detailing the patient's first visit,
    including medical history and neuropsychological evaluation information. It uses the provided
    literals for localization and formatting.

    Args:
        parsed (dict): A dictionary containing parsed patient data. Expected keys include:
            - 'patient': An object with attributes 'training_suggestion', 'date_med', and 'date_npse'.
        document (Document): A python-docx Document object where the paragraph will be added.
        literals (dict): A dictionary containing localization strings. Expected keys include:
            - 'article_v2': Article for the training suggestion sentence.
            - 'last_name': The last name of the patient.
            - 'examinee_gender': The gender of the examinee.
            - 'article_caps': Capitalized article for the first sentence.
            - 'full_name': The full name of the patient.
            - 'parents_names': The names of the patient's parents.
            - 'amka': The patient's AMKA (social security number).
            - 'unit': The unit where the patient visited.

    Returns:
        None
    """

    p1 = document.add_paragraph()

    atts_info = "" if literals["att_names"] == "" else f" και {literals['att_literal_1']} {literals['article_v3']}"
    training_suggestion = " "
    if parsed['patient'].training_suggestion is True:
        training_suggestion = f" Μετά την ολοκλήρωση της νευροψυχολογικής εκτίμησης και κατά τη διάρκεια της ανακοίνωσης των αποτελεσμάτων {literals['article_v2']} κ. {literals['last_name']}{atts_info}, συστάθηκε, συμμετοχή {literals['examinee_gender']} σε προγράμματα νοητικής ενδυνάμωσης."

    everyday_complete = ""
    if parsed['fucas'].administered:
        everyday_complete = f"αντικειμενικά από {literals['the_same_v3']}"
    if parsed['frssd'].administered:
        and_str = " και " if everyday_complete != "" else ""
        everyday_complete += f"{and_str}από {literals['att_names_with_relations']}"
    if everyday_complete != "":
        everyday_complete = f' ({everyday_complete})'

    npi_existent = f" (πληροφορίες από {literals['att_literal_1']})" if parsed['npi'].administered else ""

    medical_history_str = "όπου και ελήφθη ένα πλήρες ιατρικό και κοινωνικό ιστορικό"
    npse_diff_dates = "προκειμένου να διενεργηθεί"
    npse_same_dates = "διενεργήθηκε"
    npse_str = f" νευροψυχολογική εκτίμηση με τη χρήση συστοιχίας για την αξιολόγηση α) των νοητικών ικανοτήτων, β) της καθημερινής λειτουργικότητας{everyday_complete}, γ) της συναισθηματικής κατάστασης {literals['examinee_gender']}, καθώς επίσης και δ) των αλλαγών στη συμπεριφορά{npi_existent}.{training_suggestion}"

    date_med = parsed['patient'].date_med
    date_npse = parsed['patient'].date_npse
    # compare which date is earlier. The format is "dd/mm/yyyy"
    date_compare_result = compare_dates(date_med, date_npse)
    if date_compare_result == 0:
        p1.add_run(f"{literals['article_caps']} κ. {literals['full_name']},{literals['parents_names']}{literals['amka']} επισκέφτηκε για πρώτη φορά {literals['unit']}, στις {date_med}, {medical_history_str}. Στις {date_npse}, επισκέφτηκε την ίδια Μονάδα {npse_diff_dates}{npse_str}")
    elif date_compare_result == 2:
        p1.add_run(f"{literals['article_caps']} κ. {literals['full_name']},{literals['parents_names']}{literals['amka']} επισκέφτηκε για πρώτη φορά {literals['unit']}, στις {date_npse}, {npse_diff_dates}{npse_str} Στις {date_med}, επισκέφτηκε την ίδια Μονάδα {medical_history_str}.")
    elif date_compare_result == 1:
        p1.add_run(f"{literals['article_caps']} κ. {literals['full_name']},{literals['parents_names']}{literals['amka']} επισκέφτηκε για πρώτη φορά {literals['unit']}, στις {date_med}, {medical_history_str}. Την ίδια μέρα {npse_same_dates}{npse_str}")

    p1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
